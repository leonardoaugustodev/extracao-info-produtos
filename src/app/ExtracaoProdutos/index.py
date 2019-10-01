import glob
import docx
import json
import tqdm
import re

# Import da conexão com o mongo
from src.config import mongo as mongo


def main():

    # Conexão com o mongodb
    mongoCnxn = mongo.conectaMongo()
    banco = mongoCnxn['produtos']
    tabela = banco['teste']

    # Dropa a tabela
    tabela.delete_many({})

    # Lista de arquivos a serem processados
    listFiles = glob.glob('src/app/ExtracaoProdutos/Dados/*.docx')

    listObjects = []

    # Lista de textos para procurar nos arquivos de descrição
    # -> text: o texto a ser encontrado
    # -> column: identificador
    delimiters = [
        {
            'text': 'title',
            'column': 'Title'
        },
        {
            'text': 'description',
            'column': 'Description'
        },
        {
            'text': 'especificação técnica',
            'column': 'Specification'
        },
        {
            'text': 'cor disponível',
            'column': 'Color'
        },
        {
            'text': 'cores disponíveis',
            'column': 'Color'
        },
        {
            'text': 'lado disponível',
            'column': 'Variation'
        },
        {
            'text': 'lados disponíveis',
            'column': 'Variation'
        },
        {
            'text': 'compatível com',
            'column': 'Compatibility'
        },
        {
            'text': 'modelo disponível',
            'column': 'Design'
        },
        {
            'text': 'modelos disponíveis',
            'column': 'Design'
        },
        {
            'text': 'conteúdo da embalagem',
            'column': 'Content'
        },
        # {'text': 'importante',
        #  'column': 'Important'
        #  },
        {
            'text': 'garantia',
            'column': 'Warranty'
        },
    ]

    for file in tqdm.tqdm(listFiles):

        try:
            matches = []

            if '~' not in file:

                document = docx.Document(file)

                # Encontra o DK pai a partir do titulo do documento
                regexDK = re.search(r'\b[0-9]{5,8}\b', file, re.IGNORECASE)

                # Se não encontrar um dk, pula para o próximo arquivo
                if not regexDK:
                    continue

                dkPai = regexDK.group()

                # Se o DK contiver mais de 10 caracteres, pula para o próximo arquivo
                if len(dkPai) > 10:
                    continue

                # Encontra o paragrafo que começa o delimiter
                for (idx, p) in enumerate(document.paragraphs):

                    for delim in delimiters:
                        if(delim['text'] in p.text.lower()):
                            i = {}
                            i[delim['column']] = idx
                            matches.append(i)

                objProduct = {}
                objProduct['DK'] = dkPai

                # Para cada match na lista de matches, recupera o paragrafo
                for (idx, match) in enumerate(matches):

                    # Para o dicionário atual
                    for key, value in match.items():

                        # Encontra qual é o próximo valor
                        if(idx < len(matches) - 1):
                            nextIndex = idx + 1
                            nextValue = matches[nextIndex][next(
                                iter(matches[nextIndex]))]
                        else:
                            nextValue = len(document.paragraphs)

                        # Recupera o parágrafo
                        for para in document.paragraphs[value:nextValue]:
                            textSplited = para.text.split(':', 1)
                            text = textSplited[len(
                                textSplited) - 1].replace('•', '')

                            if text != '':

                                # Se a chave for uma das a seguir
                                if key in ('Variation', 'Compatibility', 'Specification', 'Content', 'Warranty'):

                                    textoValido = re.sub(
                                        r"[•]",  "", text, re.IGNORECASE).strip()

                                    # Se o texto tiver menos de 100 caracteres significa que é valido
                                    if(len(textoValido) < 100):

                                        if key == 'Variation':
                                            # Variation
                                            # => ok

                                            if objProduct.get(key, []) == None:
                                                oldList = []
                                            else:
                                                oldList = objProduct.get(
                                                    key, [])

                                            oldList.append(textoValido)
                                            objProduct[key] = oldList

                                        elif key == 'Compatibility':
                                            # Compatibility
                                            # { Modelo: Golf GTI, Ano: 2005 }

                                            if objProduct.get(key, []) == None:
                                                oldList = []
                                            else:
                                                oldList = objProduct.get(
                                                    key, [])

                                            reText = re.search(
                                                r"\b[0-9]{2}\b|\b[0-9]{4}\b", textoValido, re.IGNORECASE)

                                            if reText:

                                                # Se o nome do modelo houver caracteres
                                                if len(textoValido[:reText.start()]) > 0:

                                                    carro = {}
                                                    carro['Modelo'] = textoValido[:reText.start(
                                                    )]
                                                    carro['Ano'] = textoValido[reText.start(
                                                    ):]

                                                    oldList.append(carro)

                                                    objProduct[key] = oldList

                                        elif key == 'Specification':
                                            # Specification

                                            # Recupera o specification já existente
                                            if objProduct.get(key, {}) == None:
                                                specifications = []
                                            else:
                                                specifications = objProduct.get(
                                                    key, {})

                                            subKey = ''

                                            # Soquete da lâmpada
                                            if re.search(r'\b[H]+[0-9]{1,}\b', textoValido, re.IGNORECASE):

                                                subKey = 'Socket'

                                                oldSockets = objProduct.get(
                                                    key, {}).get(subKey, '')

                                                newSockets = re.findall(
                                                    r'\b[H]+[0-9]{1,}\b', textoValido, re.IGNORECASE)

                                                sockets = (oldSockets.split(
                                                    ',') + newSockets if oldSockets != '' else newSockets)

                                                specifications[subKey] = ', '.join(
                                                    sockets)

                                            # Temperatura de lâmpada
                                            elif re.search(r'\b([0-9]{4})+[K]{1}\b', textoValido, re.IGNORECASE):
                                                subKey = 'Temperature'

                                                specifications[subKey] = re.search(
                                                    r'\b([0-9]{4})+[K]{1}\b', textoValido, re.IGNORECASE).group(0)

                                            # Potencia
                                            elif re.search(r'\b[0-9]+[W]{1}\b', textoValido, re.IGNORECASE):
                                                subKey = 'Power'

                                                specifications[subKey] = textoValido[re.search(
                                                    r'\b[0-9]+[W]{1}\b', textoValido, re.IGNORECASE).start():]

                                            # Decibéis
                                            elif re.search(r'\b[0-9]+( ?[d][B])\b', textoValido, re.IGNORECASE):
                                                subKey = 'Decibels'

                                                specifications[subKey] = textoValido[re.search(
                                                    r'\b[0-9]+( ?[d][B])\b', textoValido, re.IGNORECASE).start():]

                                            # Capacidade
                                            elif re.search(r'\b[0-9]+(\s?GB|MB)\b', textoValido, re.IGNORECASE):
                                                subKey = 'Storage'

                                                specifications[subKey] = textoValido[re.search(
                                                    r'\b[0-9]+(\s?GB|MB)\b', textoValido, re.IGNORECASE).start():]

                                            if len(specifications) > len(objProduct.get(key, {}).get(subKey, '')):
                                                objProduct[key] = specifications

                                        elif key == 'Content':

                                            if text.strip() == '':
                                                continue

                                            if objProduct.get(key, []) == None:
                                                oldList = []
                                            else:
                                                oldList = objProduct.get(
                                                    key, [])

                                            oldList.append(text.strip())
                                            objProduct[key] = oldList

                                        elif key == 'Warranty':

                                            if text.strip() == '':
                                                continue

                                            if objProduct.get(key, {}) == None:
                                                oldList = {}
                                            else:
                                                oldList = objProduct.get(
                                                    key, {})

                                            # Procura onde está o texto 'Garantia de xx' no texto
                                            reText = re.search(
                                                r"^\b[garantia de ]+[0-9]{1,}\b", textoValido.lower())

                                            # Se o regex for verdadeiro, insere o texto
                                            if reText:
                                                if objProduct.get(key, None) == None:
                                                    objProduct[key] = textoValido[reText.start(
                                                    ):]

                                else:
                                    objProduct[key] = objProduct.get(
                                        key, '') + text.strip() + ('\n' if objProduct.get(
                                            key, None) else '')

                listObjects.append(objProduct)

            # Insere no banco a cada 50 registros
            if(len(listObjects) == 50):
                tabela.insert_many(listObjects)
                listObjects = []
                cont += 1

        except:
            print('Erro...')
            pass

    # Cadastra o resultado em um arquivo dados.json
    # with open('dados.json', 'w') as f:
    #     json.dump(listObjects, f)


if __name__ == '__main__':
    main()
